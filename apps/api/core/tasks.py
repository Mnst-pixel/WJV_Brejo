import subprocess
import tempfile
import zipfile
from pathlib import Path

import clamd
import pymupdf
from celery import shared_task
from django.conf import settings
from django.core.files.base import File
from django.core.files.storage import default_storage
from django.utils import timezone
from docx import Document
from PIL import Image

from .audit import record_audit
from .models import FileAsset, IngestionRun


class UnsafeDocument(Exception):
    pass


@shared_task(bind=True, autoretry_for=(ConnectionError,), retry_backoff=True, max_retries=3)
def scan_and_process_file(self, asset_id: str):
    asset = FileAsset.objects.get(pk=asset_id)
    suffix = Path(asset.original_name).suffix.lower()
    with tempfile.NamedTemporaryFile(suffix=suffix) as temp:
        with default_storage.open(asset.quarantine_key, "rb") as source:
            while chunk := source.read(1024 * 1024):
                temp.write(chunk)
        temp.flush()
        temp.seek(0)
        scanner = clamd.ClamdNetworkSocket(settings.CLAMAV_HOST, settings.CLAMAV_PORT, timeout=60)
        result = scanner.instream(temp)
        scan_result = next(iter(result.values())) if result else ("ERROR", "empty response")
        asset.scanned_at = timezone.now()
        if scan_result[0] != "OK":
            asset.scan_status = FileAsset.ScanStatus.INFECTED if scan_result[0] == "FOUND" else FileAsset.ScanStatus.ERROR
            asset.processing_status = "rejected"
            asset.metadata = {**asset.metadata, "scanner_result": scan_result[0]}
            asset.save(update_fields=["scanned_at", "scan_status", "processing_status", "metadata", "updated_at"])
            if default_storage.exists(asset.quarantine_key):
                default_storage.delete(asset.quarantine_key)
            record_audit("file.scan.rejected", actor=asset.owner, target=asset, metadata={"result": scan_result[0]})
            return

        asset.scan_status = FileAsset.ScanStatus.CLEAN
        try:
            extracted = _extract_text(temp.name, asset.mime_type)
            temp.seek(0)
            default_storage.save(asset.storage_key, File(temp, name=Path(asset.storage_key).name))
            text_key = f"processed/{asset.owner_id}/{asset.id}.txt"
            with tempfile.SpooledTemporaryFile(max_size=2 * 1024 * 1024, mode="w+b") as text_file:
                text_file.write(extracted.encode("utf-8"))
                text_file.seek(0)
                default_storage.save(text_key, File(text_file, name=f"{asset.id}.txt"))
            asset.processing_status = "processed_pending_review"
            asset.metadata = {**asset.metadata, "text_key": text_key, "extracted_characters": len(extracted)}
        except (UnsafeDocument, ValueError, OSError, RuntimeError, subprocess.SubprocessError) as exc:
            asset.processing_status = "processing_failed"
            asset.metadata = {**asset.metadata, "processing_error": type(exc).__name__}
        asset.save(update_fields=["scanned_at", "scan_status", "processing_status", "metadata", "updated_at"])
        if default_storage.exists(asset.quarantine_key):
            default_storage.delete(asset.quarantine_key)
        record_audit("file.scan.completed", actor=asset.owner, target=asset, metadata={"status": asset.processing_status})


def _extract_text(path: str, mime_type: str) -> str:
    if mime_type == "application/pdf":
        document = pymupdf.open(path)
        if document.page_count > 1000:
            raise UnsafeDocument("PDF page limit exceeded")
        text = "\n\n".join(page.get_text("text") for page in document)
    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        _validate_docx_archive(path)
        document = Document(path)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    elif mime_type.startswith("image/"):
        Image.MAX_IMAGE_PIXELS = 50_000_000
        with Image.open(path) as image:
            image.verify()
        result = subprocess.run(
            ["tesseract", path, "stdout", "-l", "por"],
            check=True,
            capture_output=True,
            timeout=90,
            text=True,
        )
        text = result.stdout
    elif mime_type == "text/plain":
        text = Path(path).read_text(encoding="utf-8", errors="strict")
    else:
        raise UnsafeDocument("Unsupported MIME")
    if len(text) > 8_000_000:
        raise UnsafeDocument("Extracted text limit exceeded")
    return text


def _validate_docx_archive(path: str):
    with zipfile.ZipFile(path) as archive:
        entries = archive.infolist()
        total_compressed = sum(item.compress_size for item in entries) or 1
        total_uncompressed = sum(item.file_size for item in entries)
        if len(entries) > 10_000 or total_uncompressed > 100 * 1024 * 1024 or total_uncompressed / total_compressed > 100:
            raise UnsafeDocument("Unsafe DOCX archive expansion")
        for item in entries:
            normalized = Path(item.filename)
            if normalized.is_absolute() or ".." in normalized.parts:
                raise UnsafeDocument("Unsafe archive path")


@shared_task
def run_ingestion(ingestion_run_id: str):
    run = IngestionRun.objects.get(pk=ingestion_run_id)
    run.status = "running"
    run.started_at = timezone.now()
    run.save(update_fields=["status", "started_at", "updated_at"])
    # Discovery adapters are deliberately source-specific and never publish here.
    run.status = "awaiting_source_adapters"
    run.report = {
        "automatic_publication": False,
        "message": "A execução foi registrada; somente adaptadores oficiais aprovados podem descobrir documentos.",
    }
    run.finished_at = timezone.now()
    run.save(update_fields=["status", "report", "finished_at", "updated_at"])
    record_audit("corpus.ingestion.finished", actor=run.requested_by, target=run, metadata={"run_type": run.run_type})
